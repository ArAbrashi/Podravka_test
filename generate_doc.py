"""
Generira Word dokument s detaljnim matematičkim opisom MILP optimizacijskog modela.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Stilovi
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================================
# NASLOV
# ============================================================
title = doc.add_heading('MILP optimizacijski model elektroenergetskog sustava tvornice', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Matematički opis mješovitog cjelobrojnog linearnog programa (MILP) '
    'za minimizaciju tjednog troška električne energije tvornice s baterijskim '
    'spremnikom, solarnom elektranom i sudjelovanjem na tržištu frekvencijske regulacije (aFRR).'
)

# ============================================================
# 1. OPIS SUSTAVA
# ============================================================
doc.add_heading('1. Opis sustava', level=1)
doc.add_paragraph(
    'Sustav se sastoji od:'
)
bullets = [
    'Tvornica s poznatom satnom potrošnjom električne energije (MW)',
    'Priključak na elektroenergetsku mrežu s ograničenom snagom povlačenja',
    'Baterijski spremnik energije (BESS) s ograničenom snagom i kapacitetom',
    'Solarna elektrana (PV) s poznatim profilom proizvodnje',
    'Mogućnost prodaje viška solarne energije u mrežu',
    'Sudjelovanje na tržištu automatske regulacije frekvencije (aFRR+ i aFRR-)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    'Optimizacija se provodi za horizont od T = 168 sati (7 dana), '
    's vremenskim korakom od 1 sata (Δt = 1 h). '
    'Budući da je Δt = 1 h, snaga u MW numerički odgovara energiji u MWh za svaki sat.'
)

# ============================================================
# 2. PARAMETRI
# ============================================================
doc.add_heading('2. Parametri modela', level=1)

table = doc.add_table(rows=1, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Simbol'
hdr[1].text = 'Naziv'
hdr[2].text = 'Vrijednost'
hdr[3].text = 'Jedinica'

params = [
    ('T', 'Broj vremenskih koraka', '168', 'h'),
    ('Δt', 'Vremenski korak', '1', 'h'),
    ('P_bat', 'Nominalna snaga baterije', '2.0', 'MW'),
    ('E_bat', 'Kapacitet baterije', '5.0', 'MWh'),
    ('SOC_init', 'Početno stanje napunjenosti', '15.0', '%'),
    ('SOC_min', 'Minimalni dopušteni SOC', '15.0', '%'),
    ('SOC_max', 'Maksimalni dopušteni SOC', '90.0', '%'),
    ('P_grid_max', 'Max snaga povlačenja iz mreže', '3.0', 'MW'),
    ('P_solar_inst', 'Instalirana snaga solarne elektrane', '2.5', 'MW'),
    ('η_chg', 'Efikasnost punjenja baterije', '0.95', '-'),
    ('η_dis', 'Efikasnost pražnjenja baterije', '0.95', '-'),
    ('n_bat_min', 'Min trajanje režima punjenja/pražnjenja', '3', 'h'),
    ('π_export', 'Cijena prodaje solarne EE u mrežu', '30.0', 'EUR/MWh'),
    ('M_deficit', 'Kazneni trošak za manjak EE', '100 000', 'EUR/MWh'),
]
for sym, name, val, unit in params:
    row = table.add_row().cells
    row[0].text = sym
    row[1].text = name
    row[2].text = val
    row[3].text = unit

doc.add_paragraph()
doc.add_paragraph('Ulazni vremenski nizovi (za svaki sat t = 0, ..., T-1):')
ts_table = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
ts_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ts_table.rows[0].cells
hdr[0].text = 'Simbol'
hdr[1].text = 'Opis'
hdr[2].text = 'Jedinica'
ts_params = [
    ('π(t)', 'Cijena električne energije na tržištu', 'EUR/MWh'),
    ('D(t)', 'Potrošnja tvornice', 'MW'),
    ('S(t)', 'Proizvodnja solarne elektrane', 'MW'),
    ('aFRR⁺(t)', 'Ponuđena snaga pozitivne regulacije', 'MW'),
    ('aFRR⁻(t)', 'Ponuđena snaga negativne regulacije', 'MW'),
]
for sym, desc, unit in ts_params:
    row = ts_table.add_row().cells
    row[0].text = sym
    row[1].text = desc
    row[2].text = unit

doc.add_paragraph()
doc.add_paragraph(
    'Solarna proizvodnja se računa iz normaliziranog profila s_norm(t) ∈ [0, 1] '
    'i instalirane snage: S(t) = s_norm(t) · P_solar_inst'
)

# ============================================================
# 3. VARIJABLE ODLUČIVANJA
# ============================================================
doc.add_heading('3. Varijable odlučivanja', level=1)

doc.add_paragraph('Za svaki sat t = 0, ..., T-1 definirane su sljedeće varijable:')

var_table = doc.add_table(rows=1, cols=5, style='Light Grid Accent 1')
var_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = var_table.rows[0].cells
hdr[0].text = 'Varijabla'
hdr[1].text = 'Tip'
hdr[2].text = 'Donja granica'
hdr[3].text = 'Gornja granica'
hdr[4].text = 'Opis'
variables = [
    ('p_grid(t)', 'Kontinuirana', '0', 'P_grid_max', 'Snaga povučena iz mreže [MW]'),
    ('p_chg(t)', 'Kontinuirana', '0', 'P_bat', 'Snaga punjenja baterije [MW]'),
    ('p_dis(t)', 'Kontinuirana', '0', 'P_bat', 'Snaga pražnjenja baterije [MW]'),
    ('SOC(t)', 'Kontinuirana', 'SOC_min', 'SOC_max', 'Stanje napunjenosti baterije [%]'),
    ('p_def(t)', 'Kontinuirana', '0', '∞', 'Manjak električne energije [MW]'),
    ('p_curt(t)', 'Kontinuirana', '0', 'S(t)', 'Curtailment solarne proizvodnje [MW]'),
    ('p_exp(t)', 'Kontinuirana', '0', 'S(t)', 'Prodaja solarne EE u mrežu [MW]'),
    ('y_chg(t)', 'Binarna', '0', '1', '1 ako baterija puni u satu t'),
    ('y_dis(t)', 'Binarna', '0', '1', '1 ako baterija prazni u satu t'),
]
for var, typ, lb, ub, desc in variables:
    row = var_table.add_row().cells
    row[0].text = var
    row[1].text = typ
    row[2].text = lb
    row[3].text = ub
    row[4].text = desc

doc.add_paragraph()
doc.add_paragraph(
    'Ukupan broj varijabli: 9 · T = 9 · 168 = 1512, '
    'od čega 336 binarnih (y_chg i y_dis).'
)

# ============================================================
# 4. FUNKCIJA CILJA
# ============================================================
doc.add_heading('4. Funkcija cilja', level=1)

doc.add_paragraph('Cilj je minimizirati ukupni tjedni trošak električne energije:')
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'min  Z = Σ [ π(t) · p_grid(t) + M_deficit · p_def(t) - π_export · p_exp(t) ]'
)
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph('gdje se sumacija provodi za t = 0, ..., T-1.')
doc.add_paragraph()
doc.add_paragraph('Objašnjenje pojedinih članova:')
items = [
    'π(t) · p_grid(t) — trošak kupnje električne energije iz mreže',
    'M_deficit · p_def(t) — kazneni trošak za nenamirenu potrošnju (penalizacija)',
    'π_export · p_exp(t) — prihod od prodaje viška solarne energije u mrežu (negativan trošak)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    'Punjenje, pražnjenje i SOC baterije nemaju direktan trošak u funkciji cilja — '
    'baterija utječe na trošak indirektno, kroz arbitražu cijena (punjenje kad je jeftino, '
    'pražnjenje kad je skupo).'
)

# ============================================================
# 5. OGRANIČENJA
# ============================================================
doc.add_heading('5. Ograničenja', level=1)

# --- 5.1 Energetska ravnoteža ---
doc.add_heading('5.1 Energetska ravnoteža', level=2)
doc.add_paragraph(
    'U svakom satu t, ukupna proizvodnja i uvoz moraju biti jednaki ukupnoj potrošnji i izvozu:'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'p_grid(t) + η_dis · p_dis(t) + S(t) + p_def(t) = D(t) + p_chg(t) + p_curt(t) + p_exp(t)'
)
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Što se može preurediti u oblik implementiran u kodu:')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'p_grid(t) + η_dis · p_dis(t) - p_chg(t) - p_curt(t) - p_exp(t) + p_def(t) = D(t) - S(t)'
)
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Objašnjenje:')
items = [
    'p_grid(t) — energija povučena iz mreže',
    'η_dis · p_dis(t) — korisna energija iz baterije (umanjena za gubitke pražnjenja)',
    'S(t) — solarna proizvodnja (besplatna energija, nije varijabla nego konstanta)',
    'p_def(t) — manjak energije (virtualni izvor s kaznenim troškom)',
    'D(t) — potrošnja tvornice (konstanta)',
    'p_chg(t) — energija usmjerena u punjenje baterije',
    'p_curt(t) — odbačena solarna energija (curtailment)',
    'p_exp(t) — solarna energija prodana u mrežu',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# --- 5.2 Dinamika SOC ---
doc.add_heading('5.2 Dinamika stanja napunjenosti (SOC)', level=2)
doc.add_paragraph(
    'SOC baterije izražen je u postocima kapaciteta (0-100%). '
    'Faktor pretvorbe iz MW·h u % kapaciteta je:'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('f = 100 / E_bat = 100 / 5.0 = 20  [%/MWh]')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Za prvi sat (t = 0):')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SOC(0) = SOC_init + η_chg · f · p_chg(0) - f · p_dis(0)')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Za ostale sate (t = 1, ..., T-1):')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SOC(t) = SOC(t-1) + η_chg · f · p_chg(t) - f · p_dis(t)')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Objašnjenje efikasnosti:')
items = [
    'η_chg · f · p_chg(t): od p_chg MW energije usmjerene u punjenje, '
    'samo η_chg udio (95%) stvarno dospije u bateriju. '
    'Npr. punjenje snagom 1 MW tijekom 1h donosi 0.95 MWh = 19% SOC.',
    'f · p_dis(t): iz baterije se uzima puna energija pražnjenja. '
    'Gubici pražnjenja (η_dis) primijenjeni su u energetskoj ravnoteži, '
    'ne u SOC dinamici — baterija gubi punu energiju, ali samo η_dis izlazi kao korisna.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# --- 5.3 SOC-ovisna max snaga pražnjenja ---
doc.add_heading('5.3 SOC-ovisno ograničenje maksimalne snage pražnjenja', level=2)
doc.add_paragraph(
    'Maksimalna snaga pražnjenja nije konstantna, već ovisi o trenutnom SOC-u. '
    'Krivulja je konkavna i definirana s 3 linearna segmenta:'
)

seg_table = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
seg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = seg_table.rows[0].cells
hdr[0].text = 'SOC raspon'
hdr[1].text = 'P_dis_max formula'
hdr[2].text = 'Opis'
segs = [
    ('0% – 50%', 'P_bat · SOC / 50', 'Linearni rast od 0 do P_bat'),
    ('50% – 85%', 'P_bat', 'Konstantno (puna snaga)'),
    ('85% – 100%', 'P_bat - 0.4·P_bat·(SOC-85)/15', 'Linearni pad do 60% P_bat'),
]
for rng, formula, desc in segs:
    row = seg_table.add_row().cells
    row[0].text = rng
    row[1].text = formula
    row[2].text = desc

doc.add_paragraph()
doc.add_paragraph(
    'Budući da je krivulja konkavna, sva tri linearna ograničenja mogu se primijeniti '
    'istovremeno bez binarnih varijabli — u svakom segmentu automatski je aktivno najstroži:'
)

doc.add_paragraph()
doc.add_heading('Segment 1:', level=3)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('p_dis(t) ≤ (P_bat / 50) · SOC(t)')
run.bold = True

doc.add_paragraph()
doc.add_heading('Segment 2:', level=3)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('p_dis(t) ≤ P_bat')
run.bold = True
doc.add_paragraph('(Osigurano gornjom granicom varijable p_dis)')

doc.add_paragraph()
doc.add_heading('Segment 3:', level=3)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('p_dis(t) + (0.4 · P_bat / 15) · SOC(t) ≤ P_bat + (0.4 · P_bat / 15) · 85')
run.bold = True

doc.add_paragraph()
doc.add_paragraph(
    'S numeričkim vrijednostima (P_bat = 2.0 MW):'
)
items = [
    'Segment 1:  p_dis(t) ≤ 0.04 · SOC(t)',
    'Segment 3:  p_dis(t) + 0.0533 · SOC(t) ≤ 6.533',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# --- 5.4 aFRR ---
doc.add_heading('5.4 Rezervacija kapaciteta za aFRR regulaciju frekvencije', level=2)
doc.add_paragraph(
    'Sudjelovanje na tržištu automatske regulacije frekvencije (aFRR) zahtijeva da baterija '
    'u svakom satu drži rezervu snage i energije za moguću aktivaciju regulacije.'
)

doc.add_heading('5.4.1 aFRR+ (pozitivna regulacija — dodatno pražnjenje)', level=3)
doc.add_paragraph(
    'Ako je baterija ponudila aFRR⁺(t) MW pozitivne regulacije, '
    'mora u svakom trenutku moći isporučiti dodatnih aFRR⁺(t) MW pražnjenja. '
    'To znači da (p_dis(t) + aFRR⁺(t)) mora zadovoljavati istu SOC-ovisnu krivulju:'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Seg. 1:  p_dis(t) ≤ (P_bat / 50) · SOC(t) - aFRR⁺(t)')
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Seg. 2:  p_dis(t) ≤ P_bat - aFRR⁺(t)')
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Seg. 3:  p_dis(t) + (0.4·P_bat/15)·SOC(t) ≤ P_bat + (0.4·P_bat/15)·85 - aFRR⁺(t)')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Energijska rezerva — SOC mora biti dovoljno visok za 1h aktivacije:')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SOC(t) ≥ SOC_min + aFRR⁺(t) · f')
run.bold = True
doc.add_paragraph()
doc.add_paragraph('Npr. za aFRR⁺ = 0.3 MW:  SOC(t) ≥ 15 + 0.3 · 20 = 21%')

doc.add_heading('5.4.2 aFRR- (negativna regulacija — dodatno punjenje)', level=3)
doc.add_paragraph(
    'Baterija mora moći primiti dodatnih aFRR⁻(t) MW punjenja:'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('p_chg(t) + aFRR⁻(t) ≤ P_bat')
run.bold = True
doc.add_paragraph()
doc.add_paragraph('tj.')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('p_chg(t) ≤ P_bat - aFRR⁻(t)')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Energijska rezerva — SOC mora biti dovoljno nizak da primi 1h dodatnog punjenja:')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SOC(t) ≤ SOC_max - aFRR⁻(t) · f')
run.bold = True
doc.add_paragraph()
doc.add_paragraph('Npr. za aFRR⁻ = 0.4 MW:  SOC(t) ≤ 90 - 0.4 · 20 = 82%')

# --- 5.5 Solar + export ---
doc.add_heading('5.5 Ograničenje solarne energije i izvoza', level=2)
doc.add_paragraph(
    'Ukupna količina solarne energije koja se odbaci (curtailment) '
    'i proda u mrežu (export) ne smije premašiti stvarnu solarnu proizvodnju:'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('p_exp(t) + p_curt(t) ≤ S(t)')
run.bold = True

doc.add_paragraph()
doc.add_paragraph(
    'Ovo osigurava da se u mrežu prodaje isključivo solarna energija — '
    'nije moguće kupiti energiju iz mreže i prodati ju po drugoj cijeni.'
)

# --- 5.6 Mutual exclusion ---
doc.add_heading('5.6 Zabrana istovremenog punjenja i pražnjenja (Big-M)', level=2)
doc.add_paragraph(
    'Baterija ne smije istovremeno puniti i prazniti. '
    'Ovo se modelira pomoću binarnih varijabli y_chg(t) i y_dis(t) '
    'te Big-M ograničenja (M = P_bat):'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('y_chg(t) + y_dis(t) ≤ 1')
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('p_chg(t) ≤ P_bat · y_chg(t)')
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('p_dis(t) ≤ P_bat · y_dis(t)')
run.bold = True

doc.add_paragraph()
doc.add_paragraph('Logika:')
items = [
    'Ako y_chg(t) = 0: punjenje je isključeno (p_chg ≤ 0 → p_chg = 0)',
    'Ako y_dis(t) = 0: pražnjenje je isključeno (p_dis ≤ 0 → p_dis = 0)',
    'y_chg + y_dis ≤ 1 osigurava da najviše jedno od dvoje može biti aktivno',
    'Mirovanje (y_chg = 0, y_dis = 0) je uvijek dopušteno',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# --- 5.7 Min trajanje ---
doc.add_heading('5.7 Minimalno trajanje režima rada baterije', level=2)
doc.add_paragraph(
    'Ako baterija počne puniti u satu t, ne smije prijeći u pražnjenje '
    'sljedećih n_bat_min - 1 sati (i obrnuto). Mirovanje je uvijek dopušteno. '
    'Ograničenje se modelira parovima nejednakosti:'
)

doc.add_paragraph()
doc.add_paragraph('Za svaki t i za svaki k ∈ {t+1, ..., min(t + n_bat_min - 1, T-1)}:')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('y_chg(t) + y_dis(k) ≤ 1     (ako puni u t, ne prazni u k)')
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('y_dis(t) + y_chg(k) ≤ 1     (ako prazni u t, ne puni u k)')
run.bold = True

doc.add_paragraph()
doc.add_paragraph(
    f'S n_bat_min = 3, to znači da nakon aktivacije punjenja ili pražnjenja '
    f'mora proći najmanje 3 sata prije promjene režima. '
    f'Unutar tog perioda baterija može ili nastaviti isti režim ili mirovati.'
)

# ============================================================
# 6. PREGLED SVIH OGRANIČENJA
# ============================================================
doc.add_heading('6. Sažetak — ukupan broj ograničenja', level=1)

doc.add_paragraph(
    'Za T = 168 sati, model generira sljedeća ograničenja:'
)

constr_table = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
constr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = constr_table.rows[0].cells
hdr[0].text = 'Ograničenje'
hdr[1].text = 'Broj po satu'
hdr[2].text = 'Ukupno'
constraints = [
    ('Energetska ravnoteža', '1', '168'),
    ('SOC dinamika', '1', '168'),
    ('SOC-ovisna snaga pražnjenja (seg. 1 i 3)', '2', '336'),
    ('aFRR+ snaga (seg. 1, 2, 3)', '3', '504'),
    ('aFRR- snaga punjenja', '1', '168'),
    ('aFRR+ energijska rezerva (SOC donji)', '1', '168'),
    ('aFRR- energijska rezerva (SOC gornji)', '1', '168'),
    ('Export + curtailment ≤ solar', '1', '168'),
    ('Mutual exclusion (y_chg + y_dis ≤ 1)', '1', '168'),
    ('Big-M punjenje (p_chg ≤ P_bat·y_chg)', '1', '168'),
    ('Big-M pražnjenje (p_dis ≤ P_bat·y_dis)', '1', '168'),
    ('Min trajanje režima (n=3)', '~4', '~664'),
]
for name, per_h, total in constraints:
    row = constr_table.add_row().cells
    row[0].text = name
    row[1].text = per_h
    row[2].text = total

row = constr_table.add_row().cells
row[0].text = 'UKUPNO'
row[1].text = ''
row[2].text = '~3018'
for cell in row:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# ============================================================
# 7. SOLVER
# ============================================================
doc.add_heading('7. Solver', level=1)
doc.add_paragraph(
    'Model se rješava pomoću HiGHS solvera (open-source, high-performance LP/MILP solver) '
    'putem Python sučelja highspy. HiGHS koristi Branch-and-Bound algoritam za '
    'cjelobrojne varijable (y_chg, y_dis) i simplex/interior-point metode za LP relaksacije.'
)

# ============================================================
# 8. IZLAZNI REZULTATI
# ============================================================
doc.add_heading('8. Izlazni rezultati', level=1)
doc.add_paragraph('Model daje sljedeće izlaze:')
items = [
    'Optimalni raspored rada sustava sat-po-sat za svih 7 dana',
    'Optimalni tjedni trošak (bez kaznenog člana za manjak)',
    'Usporedba s troškom bez baterije',
    'Ukupna energija preuzeta iz mreže, solarna proizvodnja, export, curtailment',
    'Grafički prikaz: energetska bilanca (stacked bar), SOC s aFRR granicama, cijene EE',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# SPREMI
# ============================================================
output_path = r"D:\06_Programiranje\Claude\Podravka_test\Model_opis.docx"
doc.save(output_path)
print(f"Dokument spremljen: {output_path}")
